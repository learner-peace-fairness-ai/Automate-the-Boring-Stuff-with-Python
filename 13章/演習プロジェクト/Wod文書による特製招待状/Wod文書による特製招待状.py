# "guests.txt"の招待客に招待状"invitation.docx"を作成する

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

GUEST_FILENAME = 'guests.txt'
INVITATION_FILENAME = 'invitation.docx'


def set_format(doc):
    style = doc.styles['Normal']
    style.font.name = 'Cambria'
    style.font.size = Pt(20)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


with open(GUEST_FILENAME, encoding='utf-8') as fr:
    guests = [line.rstrip() for line in fr.readlines()]

doc = Document()
set_format(doc)

for guest in guests:
    doc.add_paragraph('It would be a pleasure to have the company of')
    
    para = doc.add_paragraph()
    run = para.add_run(guest)
    run.underline = True
    
    para = doc.add_paragraph()
    run = para.add_run('at')
    run.underline = True
    para.add_run(' 11010 Memory Lane on the Evening of')
    
    doc.add_paragraph('April 1st')
    
    para = doc.add_paragraph()
    run = para.add_run('at')
    run.underline = True
    para.add_run(" 7 o'clock")

    doc.add_page_break()

doc.save(INVITATION_FILENAME)
